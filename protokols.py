from docxtpl import DocxTemplate
from docx import Document
import random
import math
import defects as df
import datetime


class CreateProt:
    '''Класс осуществляет формирование протокол контроля'''

    DATE_CONTROL = datetime.date.today().strftime('%d.%m.%Y')
    DATE_WELDING = datetime.date.today().strftime('%d.%m.%Y')

    new_defects = df.CalkDefects()


    def __init__(self,vt=False, pt=False, ut=False, rt=False, mt=False, quality_level='A',
                 number_prot='001', num_program='АСЦ-001', customer='-',
                 fio_welder='-', stigma='-', ty_opo='-', metodology='СТО Газпром 2-2.4-083-2006',
                 doc_1='-', doc_2='-', doc_3='-', doc_4='-', all_welding_types='РД',
                 kss='КСС-1', diameter=57, thickness=4, standard_size='Т+Т, 57х4', steel_grade='09Г2С', welding_method='РД', count_clicked_add_kss=0):
        self.__vt = vt
        self.__rt = rt
        self.__pt = pt
        self.__mt = mt
        self.__ut = ut

        self.__quality_level = quality_level

        self.__number_prot = number_prot
        self.__num_program = num_program
        self.__customer = customer

        self.__fio_welder = fio_welder
        self.__stigma = stigma
        self.__ty_opo = ty_opo

        self.__metodology = metodology
        self.__doc_1 = doc_1
        self.__doc_2 = doc_2
        self.__doc_3 = doc_3
        self.__doc_4 = doc_4
        self.__all_welding_types = all_welding_types

        self.__count_clicked_add_kss = count_clicked_add_kss

        self.__kss = kss
        self.__diameter = diameter
        self.__thickness = thickness
        self.__standard_size = standard_size
        self.__steel_grade = steel_grade
        self.__welding_method = welding_method


    def set_type_control(self, vt, pt, ut, rt, mt):
        '''Передаёт какой тип контроля выбран'''
        self.__vt = vt
        self.__rt = rt
        self.__pt = pt
        self.__mt = mt
        self.__ut = ut

    def set_quality_level(self, quality_level):
        '''Задаёт уровень качества'''
        if quality_level == '-':
            self.__quality_level = self.__quality_level
        else:
            self.__quality_level = quality_level

    def set_all_welding_types(self, all_welding_types):
        '''Обновляет информацию обо всех используемых типах сварки'''
        if all_welding_types:
            self.__all_welding_types = all_welding_types

    @classmethod
    def set_date_control(cls, date_control):
        '''Обновляет дату контроля, если строка не пустая.
        Значение по умолчанию - сегодняшняя дата'''
        if date_control:
            cls.DATE_CONTROL = date_control

    @classmethod
    def set_date_welding(cls, date_welding):
        '''Обновляет дату сварки, если строка не пустая.
        Значение по умолчанию - сегодняшняя дата'''
        if date_welding:
            cls.DATE_WELDING = date_welding

    def set_number_prot(self, number_prot):
        '''Обновляет номер протокола'''
        if number_prot:
            self.__number_prot = number_prot

    def set_program_for_customer(self, program, customer):
        '''Обновляет данные о заказчике'''
        if program:
            self.__num_program = program
        if customer:
            self.__customer = customer

    def set_welder_info(self, fio_welder, stigma, ty_opo):
        '''Обновляет информацию о сварщике'''
        if fio_welder:
            self.__fio_welder = fio_welder
        if stigma:
            self.__stigma = stigma
        if ty_opo:
            self.__ty_opo = ty_opo

    def set_docs(self, doc_1, doc_2, doc_3, doc_4):
        '''Обновляет документы'''
        self.__doc_1 = doc_1
        self.__doc_2 = doc_2
        self.__doc_3 = doc_3
        self.__doc_4 = doc_4

    def set_metodology(self, metodology):
        '''Обновляет методику'''
        self.__metodology = metodology

    def set_count_kss(self, count_clicked_add_kss):
        self.__count_clicked_add_kss = count_clicked_add_kss

        # Вызываю метод для добавления информации в таблицу
        self.add_row_to_table()

    def set_new_kss(self, kss, diameter, thickness, standard_size, steel_grade, welding_method):
        '''Добавляет новый образец в таблицу'''
        if kss:
            self.__kss = kss
        if diameter:
            self.__diameter = diameter
        if thickness:
            self.__thickness = thickness
        if standard_size:
            self.__standard_size = standard_size
        if steel_grade:
            self.__steel_grade = steel_grade
        if welding_method:
            self.__welding_method = welding_method


    def get_type_control(self):
        return self.__vt, self.__rt, self.__pt, self.__mt, self.__ut

    def get_quality_level(self):
        return self.__quality_level

    def get_number_prot(self):
        return self.__number_prot

    def get_program_for_customer(self):
        return self.__num_program, self.__customer

    def get_welder_info(self):
        return self.__fio_welder, self.__stigma, self.__ty_opo

    def get_docs(self):
        return self.__doc_1, self.__doc_2, self.__doc_3, self.__doc_4

    def get_metodology(self):
        return self.__metodology

    def get_count_kss(self):
        return self.__count_clicked_add_kss

    def get_new_kss(self):
        return self.__kss, self.__diameter, self.__thickness, \
               self.__standard_size, self.__steel_grade, self.__welding_method

    def get_all_welding_types(self):
        return self.__all_welding_types

    def create_context(self):
        '''Функция формирует словарь общих параметрами'''
        context = {'d_control': self.DATE_CONTROL,
                   'd_welding': self.DATE_WELDING,
                   'customer': self.__customer,
                   'program_number': self.__num_program,
                   'fio_welder': self.__fio_welder,
                   'ty_opo': self.__ty_opo,
                   'stigma': self.__stigma,
                   'number': self.__number_prot,
                   'methodology': self.__metodology,
                   'welding_method': self.__all_welding_types}
        return context

    def update_docs(self):
        '''Обновляет список документов'''
        docs = [self.__doc_1, self.__doc_2, self.__doc_3, self.__doc_4]
        return docs

    def get_current_docs(self):
        '''Возвращает текущие документы'''
        docs = self.update_docs()
        current_docs = [x for x in docs if x != '-']
        count_docs = len(current_docs)
        return count_docs, current_docs

    def get_template(self):
        '''Вызывает функции формирования предварительных протоколов
        для каждого вида контроля'''
        if self.__vt == True:
            self.get_vt_template()
        elif self.__rt == True:
            self.get_rt_template()
        elif self.__pt == True:
            self.get_pt_template()
        elif self.__mt == True:
            self.get_mt_template()
        elif self.__ut == True:
            self.get_ut_template()
        else:
            print('Не выбран метод контроля')

    def add_row_to_table(self):
        '''Записывает КСС в протоколы'''
        if self.__vt == True:
            self.add_kss_vt()
        elif self.__rt == True:
            self.add_kss_rt()
        elif self.__pt == True:
            self.add_kss_pt()
        elif self.__mt == True:
            self.add_kss_mt()
        elif self.__ut == True:
            self.add_kss_ut()
        else:
            print('Не выбран метод контроля')

    def get_vt_template(self):
        '''Формирует предварительный протокол ВИК'''
        count_docs, current_docs_list = self.get_current_docs()
        criterions = []
        context = self.create_context()
        if count_docs == 1:
            prot = DocxTemplate('pattern/VT/VT_1.docx')
            criterion = df.CalkDefects.define_criterions(doc=current_docs_list[0], vt=self.__vt)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterion})
            prot.render(context)
        elif count_docs == 2:
            prot = DocxTemplate('pattern/VT/VT_2.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, vt=self.__vt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1]})
            prot.render(context)
        elif count_docs == 3:
            prot = DocxTemplate('pattern/VT/VT_3.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, vt=self.__vt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2]})
            prot.render(context)
        else:
            prot = DocxTemplate('pattern/VT/VT_4.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, vt=self.__vt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2], 'document_4': current_docs_list[3], 'criterions_4': criterions[3]})
            prot.render(context)
        prot.save('pattern/VT/VT_with_docs/VT_pred.docx')

    def get_rt_template(self):
        '''Формирует предварительный протокол РК'''
        count_docs, current_docs_list = self.get_current_docs()
        criterions = []
        context = self.create_context()
        if count_docs == 1:
            prot = DocxTemplate('pattern/RT/RT_1.docx')
            criterion = df.CalkDefects.define_criterions(doc=current_docs_list[0], rt=self.__rt)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterion})
            prot.render(context)
        elif count_docs == 2:
            prot = DocxTemplate('pattern/RT/RT_2.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, rt=self.__rt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1]})
            prot.render(context)
        elif count_docs == 3:
            prot = DocxTemplate('pattern/RT/RT_3.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, rt=self.__rt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2]})
            prot.render(context)
        else:
            prot = DocxTemplate('pattern/RT/RT_4.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, rt=self.__rt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2], 'document_4': current_docs_list[3], 'criterions_4': criterions[3]})
            prot.render(context)
        prot.save('pattern/RT/RT_with_docs/RT_pred.docx')

    def get_pt_template(self):
        '''Формирует предварительный протокол ПВК'''
        count_docs, current_docs_list = self.get_current_docs()
        criterions = []
        context = self.create_context()
        if count_docs == 1:
            prot = DocxTemplate('pattern/PT/PT_1.docx')
            criterion = df.CalkDefects.define_criterions(doc=current_docs_list[0], pt=self.__pt)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterion})
            prot.render(context)
        elif count_docs == 2:
            prot = DocxTemplate('pattern/PT/PT_2.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, pt=self.__pt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1]})
            prot.render(context)
        elif count_docs == 3:
            prot = DocxTemplate('pattern/PT/PT_3.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, pt=self.__pt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2]})
            prot.render(context)
        else:
            prot = DocxTemplate('pattern/PT/PT_4.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, pt=self.__pt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2], 'document_4': current_docs_list[3], 'criterions_4': criterions[3]})
            prot.render(context)
        prot.save('pattern/PT/PT_with_docs/PT_pred.docx')

    def get_mt_template(self):
        '''Формирует предварительный протокол МПК'''
        count_docs, current_docs_list = self.get_current_docs()
        criterions = []
        context = self.create_context()
        if count_docs == 1:
            prot = DocxTemplate('pattern/MT/MT_1.docx')
            criterion = df.CalkDefects.define_criterions(doc=current_docs_list[0], mt=self.__mt)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterion})
            prot.render(context)
        elif count_docs == 2:
            prot = DocxTemplate('pattern/MT/MT_2.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, mt=self.__mt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1]})
            prot.render(context)
        elif count_docs == 3:
            prot = DocxTemplate('pattern/MT/MT_3.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, mt=self.__mt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2]})
            prot.render(context)
        else:
            prot = DocxTemplate('pattern/MT/MT_4.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, mt=self.__mt)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2], 'document_4': current_docs_list[3], 'criterions_4': criterions[3]})
            prot.render(context)
        prot.save('pattern/MT/MT_with_docs/MT_pred.docx')

    def get_ut_template(self):
        '''Формирует предварительный протокол УЗК'''
        count_docs, current_docs_list = self.get_current_docs()
        criterions = []
        context = self.create_context()
        if count_docs == 1:
            prot = DocxTemplate('pattern/UT/UT_1.docx')
            criterion = df.CalkDefects.define_criterions(doc=current_docs_list[0], ut=self.__ut)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterion})
            prot.render(context)
        elif count_docs == 2:
            prot = DocxTemplate('pattern/UT/UT_2.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, ut=self.__ut)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1]})
            prot.render(context)
        elif count_docs == 3:
            prot = DocxTemplate('pattern/UT/UT_3.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, ut=self.__ut)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2]})
            prot.render(context)
        else:
            prot = DocxTemplate('pattern/UT/UT_4.docx')
            for doc in current_docs_list:
                criterion = df.CalkDefects.define_criterions(doc=doc, ut=self.__ut)
                criterions.append(criterion)
            context.update({'document_1': current_docs_list[0], 'criterions_1': criterions[0], 'document_2': current_docs_list[1], 'criterions_2': criterions[1],
                            'document_3': current_docs_list[2], 'criterions_3': criterions[2], 'document_4': current_docs_list[3], 'criterions_4': criterions[3]})
            prot.render(context)
        prot.save('pattern/UT/UT_with_docs/UT_pred.docx')


    def add_general_information_for_kss(self, kss_table):
        '''Записывает основную информацию о каждом образце'''
        kss_table.rows[-1].cells[0].text = str(self.__count_clicked_add_kss)
        kss_table.rows[-1].cells[1].text = self.__kss
        kss_table.rows[-1].cells[2].text = self.DATE_WELDING
        kss_table.rows[-1].cells[3].text = self.DATE_CONTROL
        kss_table.rows[-1].cells[4].text = self.__standard_size
        kss_table.rows[-1].cells[5].text = self.__steel_grade

    def add_word_good(self, kss_table, start_column):
        '''Заполняет последние столбцы таблицы'''
        count_docs = self.get_current_docs()[0]
        start_write = start_column
        finish_column = start_write+count_docs
        for index in range(start_write, finish_column):
            kss_table.rows[-1].cells[index].text = 'годен'

    def add_tapes_rt_prot(self, count_tapes, sensitivity_rt, kss_table, count_docs, defects_list):
        '''Для заполнения количества снимков и чувствительности'''
        defects_list = defects_list
        if count_tapes != 1:
            tapes = 2
            # заполнение новых строк для плёнок в документе
            while tapes <= count_tapes:
                kss_table.add_row()
                kss_table.rows[-1].cells[6].text = str(tapes)
                kss_table.rows[-1].cells[7].text = str(sensitivity_rt)
                kss_table.rows[-1].cells[8].text = random.choice(defects_list)
                # заполнения последних столбцов
                start_write = 9
                self.add_word_good(kss_table, start_write)

                tapes = tapes+1

    def add_kss_vt(self):
        '''Добавляет новый образец в таблицу "Результаты контроля" в ВИК протокол'''
        defects_list = self.get_vt_defects()
        if self.__count_clicked_add_kss == 1:
            prot = Document('pattern/VT/VT_with_docs/VT_pred.docx')
            kss_table = prot.tables[1]
            kss_table.rows[-1].cells[1].text = self.__kss
            kss_table.rows[-1].cells[4].text = self.__standard_size
            kss_table.rows[-1].cells[5].text = self.__steel_grade
            kss_table.rows[-1].cells[6].text = random.choice(defects_list)
            prot.save('00_Reports/ВИК/'+str(self.__number_prot)+'_ВИК'+'.docx')
        else:
            prot = Document('00_Reports/ВИК/'+str(self.__number_prot)+'_ВИК'+'.docx')
            kss_table = prot.tables[1]
            kss_table.add_row()
            self.add_general_information_for_kss(kss_table)
            kss_table.rows[-1].cells[6].text = random.choice(defects_list)
            # заполнения последних столбцов
            start_write = 7
            self.add_word_good(kss_table, start_write)
            prot.save('00_Reports/ВИК/'+str(self.__number_prot)+'_ВИК'+'.docx')

    def add_kss_rt(self):
        '''Добавляет новый образец в таблицу "Результаты контроля" в РК протокол'''
        count_tapes, sensitivity_rt = self.new_defects.rt_count_tapes(self.__diameter, self.__thickness)
        defects_list = self.get_rt_defects()
        # print(defects_list)
        # для заполнения последних столбцов
        count_docs = self.get_current_docs()[0]
        if self.__count_clicked_add_kss == 1:
            prot = Document('pattern/RT/RT_with_docs/RT_pred.docx')
            kss_table = prot.tables[1]
            kss_table.rows[-1].cells[1].text = self.__kss
            kss_table.rows[-1].cells[4].text = self.__standard_size
            kss_table.rows[-1].cells[5].text = self.__steel_grade
            kss_table.rows[-1].cells[6].text = '1'
            kss_table.rows[-1].cells[7].text = str(sensitivity_rt)
            kss_table.rows[-1].cells[8].text = random.choice(defects_list)
            self.add_tapes_rt_prot(count_tapes, sensitivity_rt, kss_table, count_docs, defects_list)
            prot.save('00_Reports/РК/'+str(self.__number_prot)+'_РК'+'.docx')
        else:
            prot = Document('00_Reports/РК/'+str(self.__number_prot)+'_РК'+'.docx')
            kss_table = prot.tables[1]
            kss_table.add_row()
            self.add_general_information_for_kss(kss_table)
            kss_table.rows[-1].cells[6].text = '1'
            kss_table.rows[-1].cells[7].text = str(sensitivity_rt)
            kss_table.rows[-1].cells[8].text = random.choice(defects_list)
            # заполнения последних столбцов
            start_write = 9
            self.add_word_good(kss_table, start_write)
            # добавление снимков
            self.add_tapes_rt_prot(count_tapes, sensitivity_rt, kss_table, count_docs, defects_list)
            prot.save('00_Reports/РК/'+str(self.__number_prot)+'_РК'+'.docx')

    def add_kss_pt(self):
        '''Добавляет новый образец в таблицу "Результаты контроля" в ПВК протокол'''
        defects_list = self.get_pt_mt_defects()
        if self.__count_clicked_add_kss == 1:
            prot = Document('pattern/PT/PT_with_docs/PT_pred.docx')
            kss_table = prot.tables[1]
            kss_table.rows[-1].cells[1].text = self.__kss
            kss_table.rows[-1].cells[4].text = self.__standard_size
            kss_table.rows[-1].cells[5].text = self.__steel_grade
            kss_table.rows[-1].cells[6].text = 'II'
            kss_table.rows[-1].cells[7].text = random.choice(defects_list)
            prot.save('00_Reports/ПВК/'+str(self.__number_prot)+'_ПВК'+'.docx')
        else:
            prot = Document('00_Reports/ПВК/'+str(self.__number_prot)+'_ПВК'+'.docx')
            kss_table = prot.tables[1]
            kss_table.add_row()
            self.add_general_information_for_kss(kss_table)
            kss_table.rows[-1].cells[6].text = 'II'
            kss_table.rows[-1].cells[7].text = random.choice(defects_list)
            # заполнения последних столбцов
            start_write = 8
            self.add_word_good(kss_table, start_write)
            prot.save('00_Reports/ПВК/'+str(self.__number_prot)+'_ПВК'+'.docx')

    def add_kss_mt(self):
        '''Добавляет новый образец в таблицу "Результаты контроля" в МК протокол'''
        defects_list = self.get_pt_mt_defects()
        if self.__count_clicked_add_kss == 1:
            prot = Document('pattern/MT/MT_with_docs/MT_pred.docx')
            kss_table = prot.tables[1]
            kss_table.rows[-1].cells[1].text = self.__kss
            kss_table.rows[-1].cells[4].text = self.__standard_size
            kss_table.rows[-1].cells[5].text = self.__steel_grade
            kss_table.rows[-1].cells[8].text = random.choice(defects_list)
            prot.save('00_Reports/МПК/'+str(self.__number_prot)+'_МПК'+'.docx')
        else:
            prot = Document('00_Reports/МПК/'+str(self.__number_prot)+'_МПК'+'.docx')
            kss_table = prot.tables[1]
            kss_table.add_row()
            self.add_general_information_for_kss(kss_table)
            kss_table.rows[-1].cells[6].text = 'А'
            kss_table.rows[-1].cells[7].text = '28'
            kss_table.rows[-1].cells[8].text = random.choice(defects_list)
            # заполнения последних столбцов
            start_write = 9
            self.add_word_good(kss_table, start_write)
            prot.save('00_Reports/МПК/'+str(self.__number_prot)+'_МПК'+'.docx')

    def add_kss_ut(self):
        '''Добавляет новый образец в таблицу "Результаты контроля" в УЗК протокол'''
        defects = df.CalkDefects()
        defects_list = self.get_ut_defects()
        #print(defects_list)
        if self.__count_clicked_add_kss == 1:
            prot = Document('pattern/UT/UT_with_docs/UT_pred.docx')
            kss_table = prot.tables[1]
            kss_table.rows[-1].cells[1].text = self.__kss
            kss_table.rows[-1].cells[4].text = self.__standard_size
            kss_table.rows[-1].cells[5].text = self.__steel_grade
            # выбор преобразователя согласно методике
            piez_converter = defects.ut_select_converter(thickness=self.__thickness)
            max_eq_area = str(self.get_max_eq_area())
            kss_table.rows[-1].cells[6].text = piez_converter
            kss_table.rows[-1].cells[7].text = max_eq_area
            kss_table.rows[-1].cells[8].text = random.choice(defects_list)
            prot.save('00_Reports/УЗК/'+str(self.__number_prot)+'_УЗК'+'.docx')
        else:
            prot = Document('00_Reports/УЗК/'+str(self.__number_prot)+'_УЗК'+'.docx')
            kss_table = prot.tables[1]
            kss_table.add_row()
            self.add_general_information_for_kss(kss_table)
            piez_converter = defects.ut_select_converter(thickness=self.__thickness)
            max_eq_area = str(self.get_max_eq_area())
            kss_table.rows[-1].cells[6].text = piez_converter
            kss_table.rows[-1].cells[7].text = max_eq_area
            kss_table.rows[-1].cells[8].text = random.choice(defects_list)
            # заполнения последних столбцов
            start_write = 9
            self.add_word_good(kss_table, start_write)
            prot.save('00_Reports/УЗК/'+str(self.__number_prot)+'_УЗК'+'.docx')

    def get_max_eq_area(self):
        '''Возвращает максимальную эквивалентную площадь при УЗК'''
        list_max_eq_area = []
        defects = df.CalkDefects()
        count_docs, current_docs = self.get_current_docs()
        for doc in current_docs:
            max_eq_area = defects.ut_select_max_eq_area(doc=doc, thickness=self.__thickness, quality_level=self.__quality_level)
            list_max_eq_area.append(max_eq_area)
        max_eq_area = min(list_max_eq_area)
        return max_eq_area


    def max_possible_Aa(self):
        '''Сравнивает размер пор по всем документам и возвращает наименьший'''
        all_Aa = []
        count_docs, current_docs = self.get_current_docs()
        for doc in current_docs:
            size_Aa = self.new_defects.calk_Aa(quality_level=self.__quality_level, thickness=self.__thickness, doc=doc, welding_method=self.__welding_method)
            all_Aa.append(size_Aa)
        possible_Aa = min(all_Aa)
        print(f'Максимальной допустимый диаметр поры: {possible_Aa}')
        return possible_Aa

    def get_Fc_defect(self):
        '''Возвращает протяжённость и размер подреза'''
        count_docs, current_docs = self.get_current_docs()
        possible_Fc=[]
        if 'ГОСТ 34347-2017' not in current_docs:
            if self.__vt == True or self.__ut == True:
                for doc in current_docs:
                    defect_Fc = (f'Подрез (l={self.new_defects.calc_Fc(doc=doc)}, h=0.5, x={math.floor(random.uniform(15, (self.__diameter * math.pi - 20)))})')
                    possible_Fc.append(defect_Fc)
            elif self.__rt == True:
                for doc in current_docs:
                    defect_Fc = (f'Fc{self.new_defects.calc_Fc(doc=doc)}')
                    possible_Fc.append(defect_Fc)
            return possible_Fc
        else:
            possible_Fc.append('ДНО')
            print('Подрез не допустим')
            return possible_Fc

    def get_vt_defects(self):
        '''Варианты дефектов для ВИК'''
        all_possible_defects = ['ДНО', 'ДНО', 'ДНО']
        possible_Aa = self.max_possible_Aa()
        # генерация возможных дефектов
        variants = 3
        for i in range(variants):
            defect_Aa = (f'Пора (l={round(random.uniform(0.3, possible_Aa), 1)}, '
                         f'x={math.floor(random.uniform(15, self.__diameter * math.pi))})')
            all_possible_defects.append(defect_Aa)
        all_defects = all_possible_defects + self.get_Fc_defect()
        return all_defects

    def get_rt_defects(self):
        '''Варианты дефектов для РК'''
        all_possible_defects = ['ДНО', 'ДНО', 'ДНО']
        possible_Aa = self.max_possible_Aa()
        # генерация возможных дефектов
        variants = 3
        for i in range(variants):
            defect_Aa = (f'Aa{round(random.uniform(0.3, possible_Aa), 1)}')
            # генерация включений
            defect_Ba = (f'Ba{round(random.uniform(0.3, possible_Aa), 1)}')
            all_possible_defects.append(defect_Aa)
            all_possible_defects.append(defect_Ba)
            all_possible_defects.append('2'+defect_Aa)
            all_possible_defects.append('2'+defect_Ba)
        all_defects = all_possible_defects + self.get_Fc_defect()
        return all_defects

    def get_pt_mt_defects(self):
        '''Варианты дефектов для МПК и ПВК'''
        all_possible_defects = ['ДНО', 'ДНО']
        # генерация возможных дефектов
        variants = 3
        for i in range(variants):
            defect_Ind = (f'Индикация(l={round(random.uniform(0.3, 0.2*self.__thickness), 1)}, '
                         f'x={math.floor(random.uniform(10, (self.__diameter * math.pi - 20)))})')
            all_possible_defects.append(defect_Ind)
        return all_possible_defects

    def get_ut_defects(self):
        '''Варианты дефектов для УЗК'''
        all_possible_defects = ['ДНО', 'ДНО']
        # генерация возможных дефектов
        variants = 3
        for i in range(variants):
            if self.__thickness <= 12:
                defect_Db = (f'Несплавление(l={round(random.uniform(3, 2*self.__thickness), 1)},'
                             f'x={math.floor(random.uniform(10, (self.__diameter * math.pi - 25)))}, '
                             f'h={round(random.uniform(3, (self.__thickness - 2)), 1)})')
                all_possible_defects.append(defect_Db)
        all_defects = all_possible_defects + self.get_Fc_defect()
        return all_defects

    def print_f(self):
        print(self.context)



